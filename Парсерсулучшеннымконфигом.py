import sys
import datetime
import os
from pathlib import Path

import logging
import toml

import asyncio
import subprocess
from aiogram import Bot, Dispatcher, Router, F
from aiogram.enums import ParseMode
from aiogram.types import Message, CallbackQuery, InlineKeyboardMarkup, InlineKeyboardButton
from aiogram.client.default import DefaultBotProperties
from bs4 import BeautifulSoup
import pandas as pd
import time
import re
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from openpyxl import Workbook
from aiogram.types import FSInputFile

# ========== КОНФИГУРАЦИЯ И ЛОГГИРОВАНИЕ ==========
CONFIG_DIR = 'stat'
CURRENT_DIR = Path.cwd()
PROGRAMM_NAME = 'Парсер цен СБИС и Контур'

# Создаем папку для конфигурации если её нет
os.makedirs(CONFIG_DIR, exist_ok=True)

CONFIG_FILE_NAME = Path(CURRENT_DIR, CONFIG_DIR, 'config.toml')
LOG_FILE_NAME = Path(CURRENT_DIR, CONFIG_DIR, 'bot_log.log')

CURRENT_DATE = datetime.datetime.now().date()
CURRENT_DATE_STR = CURRENT_DATE.strftime('%d.%m.%y')

FILE_NAME_SBIS = str(Path(CURRENT_DIR, CONFIG_DIR, f'sbis_price_на_{CURRENT_DATE_STR}.xlsx'))
FILE_NAME_KONTUR = str(Path(CURRENT_DIR, CONFIG_DIR, f'kontur_price_на_{CURRENT_DATE_STR}.xlsx'))

def add_error_prefix(record):
    """Добавляет префикс ERROR только для записей с уровнем ERROR"""
    if record.levelname == "ERROR":
        record.msg = f"ERROR {record.msg}"
    else:
        record.msg = f"......{record.msg}"
    return True

logging.basicConfig(
    level=logging.INFO,
    filename=LOG_FILE_NAME,
    filemode="w",
    format='%(asctime)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logging.info("________________________________________________")
logging.info(f"*****СТАРТ программы '{PROGRAMM_NAME}'")
logging.getLogger().addFilter(add_error_prefix)

if os.path.isfile(CONFIG_FILE_NAME):
    DATA = toml.load(CONFIG_FILE_NAME)
    logging.info(f"Прочитан конфигурационный файл {CONFIG_FILE_NAME}")

    # ОТЛАДКА - добавить эти строки:
    logging.info(f"Ключи в конфиге: {list(DATA.keys())}")
    logging.info(f"Регионов СБИС в конфиге: {len(DATA.get('regions_sbis', []))}")
    logging.info(f"Регионов Контур в конфиге: {len(DATA.get('regions_kontur', []))}")
else:
    logging.error(f"Конфигурационный файл {CONFIG_FILE_NAME} не существует!")
    sys.exit()

# ========== НАСТРОЙКИ ИЗ КОНФИГА ==========
TELEGRAM_TOKEN = DATA.get('telegram', {}).get('token', '')
TELEGRAM_CHAT_ID = DATA.get('telegram', {}).get('chat_id', '')

if not TELEGRAM_TOKEN:
    logging.error("В конфигурационном файле отсутствует токен telegram!")
    sys.exit()

if not TELEGRAM_CHAT_ID:
    logging.warning("В конфигурационном файле отсутствует chat_id для отправки файлов")

async def send_file_into_chat(chat_id, doc, comment):
    """Отправляем файл в телеграм-чат"""
    try:
        logging.info(f"Начинаем отправку в чат {chat_id}")
        logging.info(f"Файл: {doc}")
        logging.info(f"Комментарий: {comment}")

        # Отправляем сообщение
        msg = await bot.send_message(chat_id=chat_id, text=comment, parse_mode='HTML')
        logging.info(f"✓ Сообщение успешно отправлено в чат {chat_id}, message_id: {msg.message_id}")

        # Отправляем файл
        rez = None
        try:
            rez = await bot.send_document(chat_id=chat_id, document=FSInputFile(doc))
            logging.info(f'✓ Файл {doc} успешно отправлен в чат {chat_id}, message_id: {rez.message_id}')
            return rez
        except Exception as e:
            logging.error(f"✗ Ошибка отправки файла: {str(e)}", exc_info=True)
            # Повторная попытка
            logging.info("Пробуем повторно отправить файл...")
            rez = await bot.send_document(chat_id=chat_id, document=FSInputFile(doc))
            logging.info(f'✓ Файл отправлен повторно, message_id: {rez.message_id}')
            return rez
    except Exception as e:
        logging.error(f"✗ КРИТИЧЕСКАЯ ошибка в send_file_into_chat: {str(e)}", exc_info=True)
        return None

# ========== НАСТРОЙКА БОТА ==========
TOKEN = TELEGRAM_TOKEN

# Создание экземпляра бота
bot = Bot(
    token=TOKEN,
    default=DefaultBotProperties(parse_mode=ParseMode.HTML)
)

# Флаг отмены парсинга
cancel_flag = False

# Кнопка отмены
cancel_button = InlineKeyboardButton(text="Отменить", callback_data="cancel_parsing")
cancel_keyboard = InlineKeyboardMarkup(inline_keyboard=[[cancel_button]])

# Создание диспетчера и маршрутизатора
dp = Dispatcher()
router = Router()
dp.include_router(router)

# Стартовая команда и кнопки
@router.message(F.text.lower() == "/start")
async def start_handler(message: Message):
    keyboard = InlineKeyboardMarkup(
        inline_keyboard=[
            [
                InlineKeyboardButton(text="СБИС", callback_data="sbis"),
                InlineKeyboardButton(text="Контур", callback_data="kontur")
            ]
        ]
    )
    await message.answer(
        "<b>Привет!</b> Извлечение цен займёт некоторое время. Выберите сайт для парсинга:",
        reply_markup=keyboard
    )

# Обработчик отмены парсинга
@router.callback_query(F.data == "cancel_parsing")
async def cancel_parsing_handler(callback_query: CallbackQuery):
    global cancel_flag
    cancel_flag = True
    # ОТВЕЧАЕМ СРАЗУ! Не ждем завершения парсинга
    await callback_query.answer("⏹ Парсинг отменяется...")
    # Дополнительно отправляем сообщение в чат
    await callback_query.message.answer("❌ Парсинг отменен пользователем.")

@router.callback_query(F.data == "sbis")
async def sbis_handler(callback_query: CallbackQuery):
    global cancel_flag
    cancel_flag = False
    await callback_query.answer("Запускаю парсинг СБИС...")
    # Отправляем кнопку "Отменить"
    await callback_query.message.answer("Парсинг СБИС начат.", reply_markup=cancel_keyboard)
    await parse_sbis(callback_query)
    await callback_query.message.answer("Парсинг СБИС завершен.")

@router.callback_query(F.data == "kontur")
async def kontur_handler(callback_query: CallbackQuery):
    global cancel_flag
    cancel_flag = False
    await callback_query.answer("Запускаю парсинг Контур...")
    # Отправляем кнопку "Отменить"
    await callback_query.message.answer("Парсинг Контур начат.", reply_markup=cancel_keyboard)
    await parse_kontur(callback_query)
    await callback_query.message.answer("Парсинг Контур завершен.")

async def parse_sbis(callback_query: CallbackQuery):
    global cancel_flag
    progress_message = await bot.send_message(callback_query.from_user.id, "СБИС: 0%")

    options = webdriver.ChromeOptions()
    options.add_argument('--headless')
    options.add_argument('--disable-gpu')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    driver = webdriver.Chrome(options=options)

    url = "https://saby.ru/tariffs?tab=ereport"
    driver.get(url)
    time.sleep(5)

    wait = WebDriverWait(driver, 20)
    all_data = []

    def safe_int(val):
        if val and str(val).isdigit():
            return int(val)
        return None

    try:
        # Загружаем регионы из конфига
        regions_to_process = DATA.get('regions_sbis', [])
        # Преобразуем в кортежи если нужно
        regions_to_process = [tuple(r) for r in regions_to_process]

        if not regions_to_process:
            logging.error("В конфиге отсутствует список регионов для СБИС!")
            await callback_query.message.answer("❌ Ошибка: список регионов не найден в конфиге")
            driver.quit()
            return

        logging.info(f"Загружено {len(regions_to_process)} регионов для СБИС из конфига")

        total = len(regions_to_process)

        for i, (region_code, region_name) in enumerate(regions_to_process):
            if cancel_flag:
                break

            progress = int((i + 1) / total * 100)
            await bot.edit_message_text(
                chat_id=callback_query.from_user.id,
                message_id=progress_message.message_id,
                text=f"СБИС: {progress}% ({i+1}/{total})"
            )

            try:
                region_url = f"https://saby.ru/tariffs?tab=ereport&region={region_code}"
                driver.get(region_url)
                WebDriverWait(driver, 15).until(
                    EC.presence_of_element_located((By.TAG_NAME, "body"))
                )
                time.sleep(3)

                driver.execute_script("window.scrollTo(0, 2500);")
                time.sleep(2)

                # ПАРСИНГ ДАННЫХ РЕГИОНА
                html = driver.page_source
                soup = BeautifulSoup(html, "html.parser")

                # ОСНОВНЫЕ ТАРИФЫ
                price_spans = soup.find_all("span", class_="billing-PriceList__priceButton")
                prices = [span.text.strip().replace(" ", "") for span in price_spans]
                filtered_prices = prices[:8] if len(prices) >= 8 else []

                # НУЛЕВКА
                null_span = soup.find("span", {"data-qa": "EOpNull"})
                null_price_raw = null_span.text.strip().replace(" ", "") if null_span else None
                null_price = safe_int(null_price_raw)

                # КОРПОРАТИВНЫЙ ТАРИФ
                corporate_prices = []
                if len(prices) >= 13:
                    corporate_prices = [
                        safe_int(prices[9]),
                        safe_int(prices[10]),
                        safe_int(prices[11]),
                        safe_int(prices[12])
                    ]

                buhta_price = None
                auth_buh_connect_price = None
                auth_buh_quarter_price = None
                auth_buh_1_199 = None
                auth_buh_200_999 = None
                auth_buh_1000_plus = None

                # ШАГ 1: Раскрываем Бухта/УПБ и извлекаем цену Бухты
                try:
                    buhta_elements = driver.find_elements(By.XPATH, "//*[contains(text(), 'Buhta') or contains(text(), 'УПБ')]")
                    for element in buhta_elements:
                        try:
                            container = element.find_element(By.XPATH, "./ancestor::div[1]")
                            container_text = container.text

                            matches = re.findall(r'(\d{1,3}\s?\d{3,4})', container_text)
                            for match in matches:
                                price_clean = match.replace(' ', '')
                                if price_clean.isdigit() and 5000 <= int(price_clean) <= 20000:
                                    buhta_price = int(price_clean)
                                    driver.execute_script("arguments[0].click();", element)
                                    time.sleep(2)
                                    break
                        except:
                            continue
                except:
                    pass

                # ШАГ 2: Уполномоченная бухгалтерия
                try:
                    auth_elements = driver.find_elements(By.XPATH, "//*[contains(text(), 'Уполномоченная бухгалтерия')]")

                    for auth_element in auth_elements:
                        try:
                            driver.execute_script("arguments[0].click();", auth_element)
                            time.sleep(3)

                            # Получаем полный текст страницы
                            page_source = driver.page_source
                            soup = BeautifulSoup(page_source, "html.parser")
                            full_text = soup.get_text()

                            # Парсим стоимость лицензии (подключение)
                            connect_match = re.search(r'Подключение[^\d]*(\d[\d\s]*)', full_text, re.IGNORECASE)
                            if connect_match:
                                connect_price_str = connect_match.group(1).replace(' ', '')
                                if connect_price_str.isdigit():
                                    auth_buh_connect_price = int(connect_price_str)

                            # Парсим за квартал (минимум)
                            quarter_match = re.search(r'(?:квартал|Квартал)[^\d]*(\d[\d\s]*)', full_text, re.IGNORECASE)
                            if not quarter_match:
                                quarter_match = re.search(r'от\s*(\d[\d\s]*)\s*[₽руб]*\s*за\s*квартал', full_text, re.IGNORECASE)
                            if quarter_match:
                                quarter_price_str = quarter_match.group(1).replace(' ', '')
                                if quarter_price_str.isdigit():
                                    auth_buh_quarter_price = int(quarter_price_str)

                            # ПАРСИНГ ЦЕН ОТЧЕТОВ
                            auth_index = full_text.find("Уполномоченная бухгалтерия")
                            if auth_index != -1:
                                auth_section = full_text[auth_index:]

                                # 1-199 (берем первые 2 цифры)
                                range_1_match = re.search(r'1[–-]199[^\d]*(\d{2,3})', auth_section)
                                if range_1_match:
                                    price_str = range_1_match.group(1)
                                    if len(price_str) >= 2:
                                        auth_buh_1_199 = int(price_str[:2])

                                # 200-999
                                range_2_match = re.search(r'200[–-]999[^\d]*(\d{2,3})', auth_section)
                                if range_2_match:
                                    auth_buh_200_999 = int(range_2_match.group(1))

                                # >1000
                                range_3_match = re.search(r'≥1\s*000\s*(\d{2,3})', auth_section)
                                if not range_3_match:
                                    range_3_match = re.search(r'≥1000\s*(\d{2,3})', auth_section)
                                if not range_3_match:
                                    range_3_match = re.search(r'>1\s*000\s*(\d{2,3})', auth_section)
                                if not range_3_match:
                                    range_3_match = re.search(r'>1000\s*(\d{2,3})', auth_section)
                                if range_3_match:
                                    auth_buh_1000_plus = int(range_3_match.group(1))

                            break

                        except:
                            continue

                except:
                    pass

                # СОБИРАЕМ ДАННЫЕ РЕГИОНА
                region_data = {
                    "Код региона": int(region_code),
                    "Название региона": region_name,
                    "Легкий_ИП": safe_int(filtered_prices[0]) if filtered_prices else None,
                    "Легкий_Бюджет": safe_int(filtered_prices[1]) if filtered_prices else None,
                    "Легкий_УСН": safe_int(filtered_prices[2]) if filtered_prices else None,
                    "Легкий_ОСНО": safe_int(filtered_prices[3]) if filtered_prices else None,
                    "Базовый_ИП": safe_int(filtered_prices[4]) if len(filtered_prices) > 4 else None,
                    "Базовый_Бюджет": safe_int(filtered_prices[5]) if len(filtered_prices) > 5 else None,
                    "Базовый_УСН": safe_int(filtered_prices[6]) if len(filtered_prices) > 6 else None,
                    "Базовый_ОСНО": safe_int(filtered_prices[7]) if len(filtered_prices) > 7 else None,
                    "Нулевка или ИП без сотрудников": null_price,
                    "ОБ (Buhta) и УПБ": buhta_price,
                    "стоимость лицензии": auth_buh_connect_price,
                    "за квартал (минимум)": auth_buh_quarter_price,
                    "1-199": auth_buh_1_199,
                    "200-999": auth_buh_200_999,
                    ">1000": auth_buh_1000_plus,
                    "5": corporate_prices[0] if corporate_prices else None,
                    "10": corporate_prices[1] if len(corporate_prices) > 1 else None,
                    "25": corporate_prices[2] if len(corporate_prices) > 2 else None,
                    "50": corporate_prices[3] if len(corporate_prices) > 3 else None,
                }

                all_data.append(region_data)

            except Exception as e:
                all_data.append({
                    "Код региона": int(region_code),
                    "Название региона": region_name,
                    "Ошибка": f"Ошибка: {str(e)}",
                })

    except Exception as e:
        pass

    # СОЗДАЕМ EXCEL ФАЙЛ С ФОРМАТИРОВАНИЕМ
    try:
        from openpyxl.styles import Font, Alignment
        from openpyxl.utils import get_column_letter

        wb = Workbook()
        ws = wb.active
        ws.title = "Цены"

        bold_font = Font(bold=True)
        center_alignment = Alignment(horizontal='center', vertical='center')

        # Заголовки
        headers_row1 = [
            "", "", "", "Легкий", "", "", "", "Базовый", "", "", "",
            "", "", "Уполномоченная бухгалтерия", "", "", "", "",
            "Корпоративный тариф", "", "", ""
        ]
        ws.append(headers_row1)

        headers_row2 = [
            "Код региона", "Название региона", "Тариф", "ИП", "Бюджет", "УСН", "ОСНО",
            "ИП", "Бюджет", "УСН", "ОСНО",
            "Нулевка или ИП без сотрудников", "ОБ (Buhta) и УПБ",
            "стоимость лицензии", "за квартал (минимум)", "1-199", "200-999", ">1000",
            "5", "10", "25", "50"
        ]
        ws.append(headers_row2)

        # Данные
        for region in all_data:
            if "Ошибка" in region:
                continue

            row_data = [
                region["Код региона"],
                region["Название региона"],
                "",
                region["Легкий_ИП"],
                region["Легкий_Бюджет"],
                region["Легкий_УСН"],
                region["Легкий_ОСНО"],
                region["Базовый_ИП"],
                region["Базовый_Бюджет"],
                region["Базовый_УСН"],
                region["Базовый_ОСНО"],
                region["Нулевка или ИП без сотрудников"],
                region["ОБ (Buhta) и УПБ"],
                region["стоимость лицензии"],
                region["за квартал (минимум)"],
                region["1-199"],
                region["200-999"],
                region[">1000"],
                region["5"],
                region["10"],
                region["25"],
                region["50"]
            ]
            ws.append(row_data)

        # Форматирование
        ws.merge_cells('D1:G1')
        ws['D1'].alignment = center_alignment
        ws['D1'].font = bold_font

        ws.merge_cells('H1:K1')
        ws['H1'].alignment = center_alignment
        ws['H1'].font = bold_font

        ws.merge_cells('M1:R1')
        ws['M1'].alignment = center_alignment
        ws['M1'].font = bold_font

        ws.merge_cells('S1:V1')
        ws['S1'].alignment = center_alignment
        ws['S1'].font = bold_font

        for col in range(1, 23):
            cell = ws.cell(row=2, column=col)
            cell.font = bold_font
            cell.alignment = center_alignment

        for row in range(3, len(all_data) + 3):
            for col in range(1, 23):
                cell = ws.cell(row=row, column=col)
                if col in [1, 2]:
                    cell.alignment = Alignment(horizontal='left', vertical='center')
                else:
                    cell.alignment = Alignment(horizontal='center', vertical='center')

        column_widths = {
            'A': 12, 'B': 20, 'C': 8, 'D': 8, 'E': 10, 'F': 8, 'G': 8,
            'H': 8, 'I': 10, 'J': 8, 'K': 8, 'L': 15, 'M': 15, 'N': 12,
            'O': 12, 'P': 8, 'Q': 8, 'R': 8, 'S': 8, 'T': 8, 'U': 8, 'V': 8
        }

        for col_letter, width in column_widths.items():
            ws.column_dimensions[col_letter].width = width

        wb.save(FILE_NAME_SBIS)

    except Exception as e:
        try:
            df = pd.DataFrame(all_data)
            df.to_excel(FILE_NAME_SBIS, index=False)
        except Exception as e2:
            pass

    driver.quit()

    await bot.edit_message_text(
        chat_id=callback_query.from_user.id,
        message_id=progress_message.message_id,
        text="✅ СБИС: Готово. Данные сохранены в saby_tariffs_filtered.xlsx"
    )

    logging.info(f"ОТЛАДКА: cancel_flag = {cancel_flag}")
    if os.path.exists(FILE_NAME_SBIS):
        logging.info(f"Файл {FILE_NAME_SBIS} создан, отправляем в чат")
        if cancel_flag:
            comment = "⚠️ Парсинг СБИС был отменен. Файл содержит неполные данные"
            logging.info("Парсинг был отменен, отправляем неполный файл")
        else:
            comment = "✅ Парсинг СБИС завершен успешно"
            logging.info("Парсинг завершен успешно")

        await send_file_into_chat(TELEGRAM_CHAT_ID, FILE_NAME_SBIS, comment)
        logging.info("Файл СБИС успешно отправлен в чат")

async def parse_kontur(callback_query: CallbackQuery):
    global cancel_flag

    # === Настройки ===
    BASE_URL = "https://www.kontur-extern.ru/price-download/77"
    DOWNLOAD_DIR = os.path.abspath("downloads")

    # === Список регионов ===
    # Загружаем регионы из конфига
    regions = DATA.get('regions_kontur', [])
    # Преобразуем в кортежи если нужно
    regions = [tuple(r) for r in regions]

    if not regions:
        logging.error("В конфиге отсутствует список регионов для Контур!")
        await callback_query.message.answer("❌ Ошибка: список регионов не найден в конфиге")
        return

    logging.info(f"Загружено {len(regions)} регионов для Контур из конфига")

    total_regions = len(regions)
    message = await callback_query.message.answer("🔄 Парсинг Контур начат...")

    # === Подготовка ===
    os.makedirs(DOWNLOAD_DIR, exist_ok=True)

    # === УЛУЧШЕННАЯ НАСТРОЙКА SELENIUM ДЛЯ HEADLESS ===
    options = webdriver.ChromeOptions()

    # Headless режим с улучшенными настройками
    options.add_argument('--headless=new')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--disable-gpu')
    options.add_argument('--window-size=1920,1080')

    # Настройки для обхода защиты и улучшения совместимости
    options.add_argument('--disable-blink-features=AutomationControlled')
    options.add_experimental_option("excludeSwitches", ["enable-automation", "enable-logging"])
    options.add_experimental_option('useAutomationExtension', False)

    # Улучшенный User-Agent
    options.add_argument('--user-agent=Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36')

    # Настройки загрузки файлов
    profile = {
        "download.default_directory": DOWNLOAD_DIR,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "plugins.always_open_pdf_externally": True,
        "safebrowsing.enabled": True,
        "profile.default_content_settings.popups": 0
    }
    options.add_experimental_option("prefs", profile)

    # Дополнительные опции для стабильности
    options.add_argument('--disable-features=VizDisplayCompositor')
    options.add_argument('--disable-software-rasterizer')
    options.add_argument('--disable-extensions')
    options.add_argument('--disable-plugins')
    options.add_argument('--disable-background-timer-throttling')
    options.add_argument('--disable-backgrounding-occluded-windows')
    options.add_argument('--disable-renderer-backgrounding')

    driver = webdriver.Chrome(options=options)

    # Улучшенное скрытие WebDriver
    driver.execute_cdp_cmd('Network.setUserAgentOverride', {
        "userAgent": 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
    })
    driver.execute_cdp_cmd('Page.addScriptToEvaluateOnNewDocument', {
        'source': '''
            Object.defineProperty(navigator, 'webdriver', {
                get: () => undefined
            });
            Object.defineProperty(navigator, 'plugins', {
                get: () => [1, 2, 3, 4, 5]
            });
            Object.defineProperty(navigator, 'languages', {
                get: () => ['ru-RU', 'ru', 'en-US', 'en']
            });
        '''
    })

    wait = WebDriverWait(driver, 30)

    # === НОВЫЕ ФУНКЦИИ ДЛЯ ИЗВЛЕЧЕНИЯ ДАННЫХ ИЗ НОВОЙ СТРУКТУРЫ ДОКУМЕНТА ===

    def convert_doc_to_docx(doc_path):
        """Конвертирует .doc в .docx используя LibreOffice"""
        try:
            docx_path = doc_path + 'x'

            try:
                subprocess.run(['libreoffice', '--version'], capture_output=True, check=True)
                libreoffice_available = True
            except:
                libreoffice_available = False

            if libreoffice_available:
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', os.path.dirname(doc_path),
                    doc_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

                if result.returncode == 0 and os.path.exists(docx_path):
                    return docx_path

            return None

        except Exception as e:
            return None

    def extract_final_price(text):
        """Извлекает итоговую цену с НДС из текста (последнее число в строке с НДС)"""
        if not text or text == "❌":
            return "❌"

        # Преобразуем в строку если нужно
        text = str(text)

        # Ищем числа в формате "X XXX,XX" или "XXXXX" - это итоговые цены с НДС
        # Они обычно в конце строки и могут быть с пробелами
        numbers = re.findall(r'(\d{1,3}(?:\s?\d{3})*(?:[.,]\d{2})?)', text)

        if numbers:
            # Берем ПОСЛЕДНЕЕ число - это итоговая стоимость с НДС
            last_number = numbers[-1]

            # Очищаем от пробелов и запятых
            clean_number = last_number.replace(' ', '').replace(',', '').replace('.', '')

            # Проверяем, что это не базовая цена (базовые обычно в 5-10 раз больше)
            if clean_number.isdigit():
                price = int(clean_number)

                # Базовая цена без НДС обычно > 100000, итоговая с НДС < 50000 для большинства тарифов
                # Но для дорогих тарифов (1+499) итоговая может быть большой
                # Поэтому проверяем по контексту позже

                return price

        return "❌"

    def extract_optimal_plus_from_table(table, results):
        """Извлекает данные из таблицы Оптимальный плюс"""
        try:
            rows = list(table.rows)

            # Ищем строки с "Оптимальный плюс" и "1 год"
            for i, row in enumerate(rows):
                row_text = [cell.text.strip() for cell in row.cells]
                row_lower = ' '.join(row_text).lower()

                if "оптимальный плюс" in row_lower and "1 год" in row_lower:
                    # Проверяем следующие строки для разных категорий
                    for j in range(i, min(i+8, len(rows))):
                        check_row = rows[j]
                        check_text = ' '.join([c.text.lower() for c in check_row.cells])

                        cells = check_row.cells
                        if len(cells) >= 8:

                            # Ищем ячейку с итоговой стоимостью (последняя колонка)
                            final_price_cell = cells[-1].text

                            # Определяем категорию по тексту
                            if "ип" in check_text:
                                if "усн" in check_text or "специальная" in check_text:
                                    # Для ИП УСН итоговая цена 6 500,00
                                    price = extract_final_price(final_price_cell)
                                    if price and price != "❌" and 5000 < price < 10000:
                                        results['ip_usn'] = price

                                elif "общая" in check_text or "осно" in check_text or "смешанная" in check_text:
                                    # Для ИП ОСНО итоговая цена 9 500,00
                                    price = extract_final_price(final_price_cell)
                                    if price and price != "❌" and 8000 < price < 12000:
                                        results['ip_osno'] = price

                            elif "юл" in check_text:
                                if "усн" in check_text or "специальная" in check_text:
                                    # Для ЮЛ УСН итоговая цена 9 500,00
                                    price = extract_final_price(final_price_cell)
                                    if price and price != "❌" and 8000 < price < 12000:
                                        results['ul_usn'] = price

                                elif "общая" in check_text or "осно" in check_text or "смешанная" in check_text:
                                    # Для ЮЛ ОСНО итоговая цена 12 500,00
                                    price = extract_final_price(final_price_cell)
                                    if price and price != "❌" and 10000 < price < 15000:
                                        results['ul_osno'] = price
        except Exception as e:
            pass

    def extract_budget_plus_from_table(table, results):
        """Извлекает данные из таблицы Бюджетник плюс"""
        try:
            rows = list(table.rows)

            for row in rows:
                cells = row.cells
                if len(cells) >= 6:
                    row_text = ' '.join([c.text.lower() for c in cells])

                    # Ищем строку с "Бюджетник плюс" и "1 год"
                    if "бюджетник плюс" in row_text and "1 год" in row_text:
                        # Итоговая стоимость в последней колонке
                        final_price = extract_final_price(cells[-1].text)
                        if final_price and final_price != "❌" and 5000 < final_price < 10000:
                            results['budget_plus'] = final_price

                    # Ищем строку с "Бюджетник Максимальный" и "1 год"
                    elif "бюджетник максимальный" in row_text and "1 год" in row_text:
                        final_price = extract_final_price(cells[-1].text)
                        if final_price and final_price != "❌" and 10000 < final_price < 20000:
                            results['budget'] = final_price
        except Exception as e:
            pass

    def extract_common_tariffs_from_table(table, results, common_keys):
        """Извлекает данные из таблицы Общий и Общий плюс"""
        try:
            rows = list(table.rows)

            for i, row in enumerate(rows):
                cells = row.cells
                if len(cells) >= 4:
                    row_text = ' '.join([c.text.lower() for c in cells])

                    # Ищем строки с "Общий" (без плюс) для первого года
                    if "общий" in row_text and "плюс" not in row_text and "1 год" in row_text:
                        # Проверяем все ключи
                        for key in common_keys:
                            key_lower = key.lower().replace('+', '').replace(' ', '')
                            if key_lower in row_text.replace(' ', '').replace('+', ''):
                                # Итоговая стоимость в последней колонке
                                final_price = extract_final_price(cells[-1].text)
                                if final_price and final_price != "❌":
                                    # Проверяем соответствие ожидаемым значениям
                                    expected_ranges = {
                                        "1+4": (10000, 20000),      # 14 500
                                        "1+9": (15000, 25000),      # 18 900
                                        "1+19": (20000, 35000),     # 28 900
                                        "1+49": (40000, 70000),     # 58 500
                                        "1+99": (70000, 100000),    # 89 000
                                        "1+199": (150000, 200000),  # 168 500
                                        "1+499": (300000, 350000)   # 319 600
                                    }
                                    if key in expected_ranges:
                                        min_val, max_val = expected_ranges[key]
                                        if min_val <= final_price <= max_val:
                                            results['common'][key] = final_price
                                break

                    # Ищем строки с "Общий плюс" для первого года
                    elif "общий плюс" in row_text and "1 год" in row_text:
                        for key in common_keys:
                            key_lower = key.lower().replace('+', '').replace(' ', '')
                            if key_lower in row_text.replace(' ', '').replace('+', ''):
                                final_price = extract_final_price(cells[-1].text)
                                if final_price and final_price != "❌":
                                    # Ожидаемые диапазоны для Общий плюс
                                    expected_ranges = {
                                        "1+4": (20000, 30000),      # 24 200
                                        "1+9": (25000, 35000),      # 30 800
                                        "1+19": (35000, 50000),     # 42 400
                                        "1+49": (80000, 100000),    # 90 900
                                        "1+99": (130000, 160000),   # 145 400
                                        "1+199": (250000, 300000),  # 269 500
                                        "1+499": (400000, 450000)   # 418 900
                                    }
                                    if key in expected_ranges:
                                        min_val, max_val = expected_ranges[key]
                                        if min_val <= final_price <= max_val:
                                            results['common_plus'][key] = final_price
                                break
        except Exception as e:
            pass

    def extract_prices_universal(filepath):
        """Универсальное извлечение цен из Word документов"""
        try:
            file_ext = os.path.splitext(filepath)[1].lower()

            if file_ext == '.docx':
                return extract_from_docx_by_structure(filepath)
            elif file_ext == '.doc':
                converted_path = convert_doc_to_docx(filepath)
                if converted_path:
                    return extract_from_docx_by_structure(converted_path)

            return ["❌"] * 22

        except Exception as e:
            return ["❌"] * 22

    def convert_doc_to_docx(doc_path):
        """Конвертирует .doc в .docx используя LibreOffice"""
        try:
            docx_path = doc_path + 'x'

            try:
                subprocess.run(['libreoffice', '--version'], capture_output=True, check=True)
                libreoffice_available = True
            except:
                libreoffice_available = False

            if libreoffice_available:
                cmd = [
                    'libreoffice', '--headless', '--convert-to', 'docx',
                    '--outdir', os.path.dirname(doc_path),
                    doc_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)

                if result.returncode == 0 and os.path.exists(docx_path):
                    return docx_path

            return None
        except Exception as e:
            print(f"Ошибка конвертации: {e}")
            return None

    def extract_number_from_cell(text):
        """Извлекает число из ячейки таблицы"""
        if not text:
            return "❌"

        # Ищем число (с пробелами или без)
        text = str(text)
        # Убираем пробелы и заменяем запятую на точку
        cleaned = text.replace(' ', '').replace(',', '.').replace('–', '').strip()

        # Ищем число в формате XXXX.XX или XXXXX
        match = re.search(r'(\d+(?:\.\d+)?)', cleaned)
        if match:
            num_str = match.group(1)
            if '.' in num_str:
                num_str = num_str.split('.')[0]
            if num_str.isdigit():
                return int(num_str)

        return "❌"

    def extract_from_docx_by_structure(filepath):
        """Извлечение данных по структуре документа"""
        try:
            from docx import Document
            doc = Document(filepath)

            # Инициализация результатов
            ip_usn = "❌"
            ip_osno = "❌"
            ul_usn = "❌"
            ul_osno = "❌"
            budget_plus = "❌"
            budget = "❌"
            common_prices = ["❌"] * 7  # 1+4 до 1+499
            common_plus_prices = ["❌"] * 7  # 1+4 плюс до 1+499 плюс

            # Получаем все таблицы
            tables = list(doc.tables)

            # ===== ТАБЛИЦА 1: Оптимальный плюс =====
            if len(tables) >= 1:
                table = tables[0]
                rows = list(table.rows)

                # Ищем строки с "Оптимальный плюс" и "1 год"
                for i, row in enumerate(rows):
                    cells = row.cells
                    if len(cells) >= 8:
                        # Получаем текст всех ячеек для анализа
                        row_text = ' '.join([c.text.lower() for c in cells])

                        # Проверяем, что это строка с данными (не заголовок)
                        if "оптимальный плюс" in row_text and "1 год" in row_text:
                            # Определяем категорию
                            if "ип" in row_text:
                                if "усн" in row_text or "специальная" in row_text:
                                    # ИП УСН - берем цену из последней ячейки
                                    price = extract_number_from_cell(cells[-1].text)
                                    if price != "❌":
                                        ip_usn = price
                                elif "общая" in row_text or "осно" in row_text or "смешанная" in row_text:
                                    # ИП ОСНО
                                    price = extract_number_from_cell(cells[-1].text)
                                    if price != "❌":
                                        ip_osno = price
                            elif "юл" in row_text:
                                if "усн" in row_text or "специальная" in row_text:
                                    # ЮЛ УСН
                                    price = extract_number_from_cell(cells[-1].text)
                                    if price != "❌":
                                        ul_usn = price
                                elif "общая" in row_text or "осно" in row_text or "смешанная" in row_text:
                                    # ЮЛ ОСНО
                                    price = extract_number_from_cell(cells[-1].text)
                                    if price != "❌":
                                        ul_osno = price

            # ===== ТАБЛИЦА 3: Бюджетник (индекс 2) =====
            if len(tables) >= 3:
                table = tables[2]  # Третья таблица (индекс 2)
                rows = list(table.rows)

                # Сбрасываем найденные значения
                found_budget_plus = False
                found_budget_normal = False

                for row in rows:
                    cells = row.cells
                    if len(cells) >= 6:
                        row_text = ' '.join([c.text.lower() for c in cells])

                        # Пропускаем строки с "Максимальный" - они нам не нужны
                        if "максимальный" in row_text:
                            continue

                        # Ищем "Бюджетник плюс" (срок 1 год)
                        if "бюджетник плюс" in row_text and "1 год" in row_text and not found_budget_plus:
                            price = extract_number_from_cell(cells[-1].text)
                            if price != "❌":
                                budget_plus = price
                                found_budget_plus = True

                        # Ищем обычный "Бюджетник" (без "плюс" и без "максимальный") со сроком 1 год
                        elif "бюджетник" in row_text and "плюс" not in row_text and "1 год" in row_text and not found_budget_normal:
                            # Проверяем, что это действительно обычный бюджетник
                            if not any(word in row_text for word in ["максимальный", "плюс"]):
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    budget = price
                                    found_budget_normal = True

            # ===== ТАБЛИЦА 5: Общий (индекс 4) =====
            if len(tables) >= 5:
                table = tables[4]  # Пятая таблица (индекс 4)
                rows = list(table.rows)

                common_index = 0
                for row in rows:
                    cells = row.cells
                    if len(cells) >= 7:
                        row_text = ' '.join([c.text.lower() for c in cells])

                        # Ищем строки с "Общий" (без плюс) и "1 год"
                        if "общий" in row_text and "плюс" not in row_text and "1 год" in row_text:
                            # Определяем количество абонентов
                            if "1+4" in row_text and common_index == 0:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_prices[0] = price
                                    common_index += 1
                            elif "1+9" in row_text and common_index <= 1:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_prices[1] = price
                                    common_index += 1
                            elif "1+19" in row_text and common_index <= 2:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_prices[2] = price
                                    common_index += 1
                            elif "1+49" in row_text and common_index <= 3:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_prices[3] = price
                                    common_index += 1
                            elif "1+99" in row_text and common_index <= 4:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_prices[4] = price
                                    common_index += 1
                            elif "1+199" in row_text and common_index <= 5:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_prices[5] = price
                                    common_index += 1
                            elif "1+499" in row_text and common_index <= 6:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_prices[6] = price
                                    common_index += 1

            # ===== ТАБЛИЦА 6: Общий плюс (индекс 5) =====
            if len(tables) >= 6:
                table = tables[5]  # Шестая таблица (индекс 5)
                rows = list(table.rows)

                common_plus_index = 0
                for row in rows:
                    cells = row.cells
                    if len(cells) >= 7:
                        row_text = ' '.join([c.text.lower() for c in cells])

                        # Ищем строки с "Общий плюс" и "1 год"
                        if "общий плюс" in row_text and "1 год" in row_text:
                            # Определяем количество абонентов
                            if "1+4" in row_text and common_plus_index == 0:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_plus_prices[0] = price
                                    common_plus_index += 1
                            elif "1+9" in row_text and common_plus_index <= 1:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_plus_prices[1] = price
                                    common_plus_index += 1
                            elif "1+19" in row_text and common_plus_index <= 2:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_plus_prices[2] = price
                                    common_plus_index += 1
                            elif "1+49" in row_text and common_plus_index <= 3:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_plus_prices[3] = price
                                    common_plus_index += 1
                            elif "1+99" in row_text and common_plus_index <= 4:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_plus_prices[4] = price
                                    common_plus_index += 1
                            elif "1+199" in row_text and common_plus_index <= 5:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_plus_prices[5] = price
                                    common_plus_index += 1
                            elif "1+499" in row_text and common_plus_index <= 6:
                                price = extract_number_from_cell(cells[-1].text)
                                if price != "❌":
                                    common_plus_prices[6] = price
                                    common_plus_index += 1

            # Формируем результат в нужном порядке
            result = [
                ip_usn,           # колонка 3: ИП (УСН)
                ip_osno,          # колонка 4: ИП (ОСНО)
                ul_usn,           # колонка 5: ЮЛ (УСН)
                ul_osno,          # колонка 6: ЮЛ (ОСНО)
                budget_plus,      # колонка 7: Бюджетник плюс
                budget,           # колонка 8: Обычный Бюджетник (или ❌ если нет)
                common_prices[0], # колонка 9: 1+4
                common_prices[1], # колонка 10: 1+9
                common_prices[2], # колонка 11: 1+19
                common_prices[3], # колонка 12: 1+49
                common_prices[4], # колонка 13: 1+99
                common_prices[5], # колонка 14: 1+199
                common_prices[6], # колонка 15: 1+499
                common_plus_prices[0], # колонка 16: 1+4 плюс
                common_plus_prices[1], # колонка 17: 1+9 плюс
                common_plus_prices[2], # колонка 18: 1+19 плюс
                common_plus_prices[3], # колонка 19: 1+49 плюс
                common_plus_prices[4], # колонка 20: 1+99 плюс
                common_plus_prices[5], # колонка 21: 1+199 плюс
                common_plus_prices[6]  # колонка 22: 1+499 плюс
            ]

            return result

        except Exception as e:
            import traceback
            traceback.print_exc()
            return ["❌"] * 22

    # === СТАРЫЕ ФУНКЦИИ ДЛЯ PDF (ОСТАВЛЯЕМ БЕЗ ИЗМЕНЕНИЙ) ===

    def extract_text_from_pdf(pdf_path):
        """Извлекает текст из PDF файла"""
        try:
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += page_text
                return text
        except Exception as e:
            return ""

    def extract_all_null_prices(pdf_path):
        """
        Извлекает итоговую стоимость с НДС для Нулевой отчетности по всем регионам
        """
        import PyPDF2
        import re

        try:
            null_reporting_data = {}

            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # Страницы с Нулевой отчетностью (49-54 в документе = индексы 48-53)
                for page_num in range(48, 54):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    lines = text.split('\n')

                    for line in lines:
                        line_clean = line.strip()

                        # Ищем строки с "Право использования ПО"
                        if 'Право использования ПО' in line_clean and len(line_clean) >= 2 and line_clean[:2].isdigit():
                            region_code = line_clean[:2]

                            # Ищем паттерн: "– число" (итоговая стоимость после тире)
                            # Формат: "... – 2 200,00 ..."
                            match = re.search(r'–\s+([\d\s,]+)', line_clean)
                            if match:
                                price_str = match.group(1).strip()
                                # Убираем пробелы и заменяем запятую на точку
                                price_str = price_str.replace(' ', '').replace(',', '.')

                                try:
                                    price = float(price_str)
                                    null_reporting_data[region_code] = price
                                except ValueError:
                                    continue

            return null_reporting_data

        except Exception as e:
            print(f"Ошибка при парсинге Нулевой отчетности: {e}")
            import traceback
            traceback.print_exc()
            return {}

    def extract_all_tax_representative_prices(pdf_path):
        """Извлекает все цены налогового представителя из PDF с учетом регрессивных шкал"""
        text = extract_text_from_pdf(pdf_path)
        if not text:
            return {}

        regression_zones = extract_regression_zones(text)

        if not regression_zones:
            pass

        lines = text.split('\n')
        prices_dict = {}

        # Список всех настоящих кодов регионов
        real_region_codes = [str(i).zfill(2) for i in range(1, 96)]
        real_region_codes += ['77', '78', '79', '83', '86', '87', '89', '90', '91', '92', '93', '94', '95', '99']

        # Объединяем строки для каждого региона
        current_region = ""
        combined_text = ""

        for line in lines:
            line_clean = line.strip()
            if not line_clean:
                continue

            # Строгая проверка: строка должна начинаться с настоящего кода региона и содержать название
            is_region_line = False
            for region_code in real_region_codes:
                if (line_clean.startswith(region_code + ' ') and
                    len(line_clean) > 10 and
                    any(char.isalpha() for char in line_clean[3:10])):
                    is_region_line = True
                    break

            if is_region_line:
                if current_region and combined_text:
                    process_tax_region_with_zones(current_region, combined_text, prices_dict, real_region_codes, regression_zones)

                current_region = line_clean.split()[0] if line_clean.split() else ""
                combined_text = line_clean
            else:
                if current_region:
                    combined_text += " " + line_clean

        if current_region and combined_text:
            process_tax_region_with_zones(current_region, combined_text, prices_dict, real_region_codes, regression_zones)

        return prices_dict

    def extract_regression_zones(text):
        """Извлекает данные регрессивных шкал из текста PDF"""
        zones = {}

        lines = text.split('\n')

        # Создаем структуры для всех зон
        all_zones = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12']
        for zone in all_zones:
            zones[zone] = {}

        zone_headers = ['1', '2', '3', '5', '6', '7', '8', '9', '11', '12']

        for i, line in enumerate(lines):
            line_clean = line.strip()

            if "До 199" in line_clean or "До 192" in line_clean:
                all_numbers = re.findall(r'\b(\d{2,3})\b', line_clean)
                prices = all_numbers[1:] if len(all_numbers) > 1 else []

                if len(prices) >= len(zone_headers):
                    for j, price_str in enumerate(prices):
                        if j < len(zone_headers):
                            # ОЧИЩАЕМ от нецифровых символов
                            clean_str = re.sub(r'[^\d]', '', price_str)
                            if clean_str.isdigit():
                                zone_num = zone_headers[j]
                                zones[zone_num]["до_199"] = int(clean_str)

            elif "От 200 до 499" in line_clean:
                parts = line_clean.split("499")
                if len(parts) > 1:
                    prices_part = parts[1]
                    prices = re.findall(r'\b(\d{2,3})\b', prices_part)

                    if len(prices) >= len(zone_headers):
                        for j, price_str in enumerate(prices):
                            if j < len(zone_headers):
                                # ОЧИЩАЕМ от нецифровых символов
                                clean_str = re.sub(r'[^\d]', '', price_str)
                                if clean_str.isdigit():
                                    zone_num = zone_headers[j]
                                    zones[zone_num]["от_200_до_499"] = int(clean_str)

            elif "От 500 до 999" in line_clean:
                parts = line_clean.split("999")
                if len(parts) > 1:
                    prices_part = parts[1]
                    prices = re.findall(r'\b(\d{2,3})\b', prices_part)

                    if len(prices) >= len(zone_headers):
                        for j, price_str in enumerate(prices):
                            if j < len(zone_headers):
                                # ОЧИЩАЕМ от нецифровых символов
                                clean_str = re.sub(r'[^\d]', '', price_str)
                                if clean_str.isdigit():
                                    zone_num = zone_headers[j]
                                    zones[zone_num]["от_500_до_999"] = int(clean_str)

            elif "От 1000 до 1999" in line_clean:
                parts = line_clean.split("1999")
                if len(parts) > 1:
                    prices_part = parts[1]
                    prices = re.findall(r'\b(\d{2,3})\b', prices_part)

                    if len(prices) >= len(zone_headers):
                        for j, price_str in enumerate(prices):
                            if j < len(zone_headers):
                                # ОЧИЩАЕМ от нецифровых символов
                                clean_str = re.sub(r'[^\d]', '', price_str)
                                if clean_str.isdigit():
                                    zone_num = zone_headers[j]
                                    zones[zone_num]["от_1000_до_1999"] = int(clean_str)

            elif "От 2000" in line_clean and "От 2000 до" not in line_clean:
                parts = line_clean.split("2000")
                if len(parts) > 1:
                    prices_part = parts[1]
                    prices = re.findall(r'\b(\d{2,3})\b', prices_part)

                    if len(prices) >= len(zone_headers):
                        for j, price_str in enumerate(prices):
                            if j < len(zone_headers):
                                # ОЧИЩАЕМ от нецифровых символов
                                clean_str = re.sub(r'[^\d]', '', price_str)
                                if clean_str.isdigit():
                                    zone_num = zone_headers[j]
                                    zones[zone_num]["от_2000"] = int(clean_str)

        # ПАРСИМ ДАННЫЕ ДЛЯ ЗОН 4 И 10 ОТДЕЛЬНО (ИЗ ДРУГОЙ ТАБЛИЦЫ)
        for i, line in enumerate(lines):
            line_clean = line.strip()

            # Ищем данные для зон 4 и 10 с их специфичными диапазонами
            if "До 349" in line_clean:
                all_numbers = re.findall(r'\b(\d{2,3})\b', line_clean)
                prices = all_numbers[1:] if len(all_numbers) > 1 else []  # Исключаем 349
                if len(prices) >= 2:
                    # ОЧИЩАЕМ от нецифровых символов
                    clean_price1 = re.sub(r'[^\d]', '', prices[0])
                    clean_price2 = re.sub(r'[^\d]', '', prices[1])
                    if clean_price1.isdigit():
                        zones["4"]["до_349"] = int(clean_price1)
                    if clean_price2.isdigit():
                        zones["10"]["до_349"] = int(clean_price2)

            elif "От 350 до 599" in line_clean:
                parts = line_clean.split("599")
                if len(parts) > 1:
                    prices_part = parts[1]
                    prices = re.findall(r'\b(\d{2,3})\b', prices_part)
                    if len(prices) >= 2:
                        # ОЧИЩАЕМ от нецифровых символов
                        clean_price1 = re.sub(r'[^\d]', '', prices[0])
                        clean_price2 = re.sub(r'[^\d]', '', prices[1])
                        if clean_price1.isdigit():
                            zones["4"]["от_350_до_599"] = int(clean_price1)
                        if clean_price2.isdigit():
                            zones["10"]["от_350_до_599"] = int(clean_price2)

            elif "От 600 до 999" in line_clean:
                parts = line_clean.split("999")
                if len(parts) > 1:
                    prices_part = parts[1]
                    prices = re.findall(r'\b(\d{2,3})\b', prices_part)
                    if len(prices) >= 2:
                        # ОЧИЩАЕМ от нецифровых символов
                        clean_price1 = re.sub(r'[^\d]', '', prices[0])
                        clean_price2 = re.sub(r'[^\d]', '', prices[1])
                        if clean_price1.isdigit():
                            zones["4"]["от_600_до_999"] = int(clean_price1)
                        if clean_price2.isdigit():
                            zones["10"]["от_600_до_999"] = int(clean_price2)

            # Строка "От 1000" для зон 4 и 10 (у них только один диапазон "от 1000")
            elif "От 1000" in line_clean:
                parts = line_clean.split()
                for idx, part in enumerate(parts):
                    if part == "1000" and idx + 2 < len(parts):
                        # ОЧИЩАЕМ от нецифровых символов
                        clean_price1 = re.sub(r'[^\d]', '', parts[idx + 1])
                        clean_price2 = re.sub(r'[^\d]', '', parts[idx + 2])
                        if clean_price1.isdigit():
                            zones["4"]["от_1000"] = int(clean_price1)
                        if clean_price2.isdigit():
                            zones["10"]["от_1000"] = int(clean_price2)
                        break

        return zones

    def process_tax_region_with_zones(region_code, text, prices_dict, real_region_codes, regression_zones):
        """Обрабатывает один регион с учетом регрессивных шкал"""
        if region_code not in real_region_codes:
            return

        # ВАЖНО: Если регион уже обработан, не перезаписываем!
        if region_code in prices_dict:
            return

        zone_match = re.search(r'(\d{1,2})(?=\s+\d+\s+\d+\s+\d+\s+\d+)', text)
        zone_number = None

        if zone_match:
            zone_number = zone_match.group(1)
        else:
            numbers = re.findall(r'\b(\d{1,2})\b', text)
            for num in numbers:
                if num in ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12']:
                    zone_number = num
                    break

        tax_data = {
            "zone": zone_number,
            "base_price": None,
            "regression_prices": {}
        }

        # Ищем паттерн: текст между "Право" и "Услуги"
        right_pattern = r'Право\s+(.*?)\s+Услуги'
        right_match = re.search(right_pattern, text)

        if right_match:
            right_text = right_match.group(1)
            # Находим все цены
            prices = re.findall(r'(\d[\d\s]*,\d+)', right_text)

            # Четвёртая цена (индекс 3) = итоговая за 12 месяцев
            if len(prices) >= 4:
                tax_price_str = prices[3].replace(' ', '').replace(',', '.')

                try:
                    tax_price = float(tax_price_str)

                    # ФИЛЬТР: Базовый имеет цены в диапазоне 6500-17000
                    if 6500 <= tax_price <= 17000:
                        tax_data["base_price"] = tax_price

                        if zone_number and zone_number in regression_zones:
                            tax_data["regression_prices"] = regression_zones[zone_number]

                        prices_dict[region_code] = tax_data
                except ValueError:
                    pass

        return

    def extract_all_start_online_prices(pdf_path):
        """Извлекает все цены Стартовый онлайн из PDF"""
        text = extract_text_from_pdf(pdf_path)
        if not text:
            return {}

        lines = text.split('\n')
        prices_dict = {}

        current_region = ""
        current_text = ""

        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue

            if re.match(r'^\d{2}', line):
                if current_region and current_text:
                    process_region_for_start_online_improved(current_region, current_text, prices_dict)

                current_region = line.split()[0] if line.split() else ""
                current_text = line
            else:
                if current_region:
                    current_text += " " + line

        if current_region and current_text:
            process_region_for_start_online_improved(current_region, current_text, prices_dict)

        return prices_dict

    def process_region_for_start_online_improved(region_code, text, prices_dict):
        """Обрабатывает текст региона для извлечения цен Стартовый онлайн"""

        # Ищем все пары чисел в формате "число1 – число2" где число2 - итоговая цена
        pattern = r'(\d[\d\s,\.]*)\s*–\s*(\d[\d\s,\.]+)'
        matches = re.findall(pattern, text)

        prices = []

        for base_price, final_price in matches:
            # Очищаем итоговую цену (второе число после тире)
            clean_price = final_price.replace(' ', '').replace(',', '').replace('\xa0', '').strip()

            # Число приходит с копейками: "4 800,00" -> "480000"
            # Делим на 100 чтобы получить правильную цену
            if clean_price.isdigit() and len(clean_price) >= 5:
                price = int(clean_price) // 100

                if 3000 <= price <= 20000 and price != int(region_code):
                    prices.append(price)

        # НЕ удаляем дубликаты! Нам нужны все 4 цены для 4 категорий
        if len(prices) >= 4:
            prices_dict[region_code] = prices[:4]
        else:
            alternative_prices = extract_start_online_alternative_improved(text, region_code)
            if alternative_prices and len(alternative_prices) >= 4:
                prices_dict[region_code] = alternative_prices

    def extract_start_online_alternative_improved(text, region_code):
        """Альтернативный метод извлечения цен Стартовый онлайн"""
        spaced_prices = re.findall(r'(\d{1,2}\s?\d{3})', text)
        if spaced_prices:
            prices = []
            for price_str in spaced_prices:
                clean_price = int(price_str.replace(' ', ''))
                if 3000 <= clean_price <= 20000 and clean_price != int(region_code):
                    prices.append(clean_price)
                    if len(prices) >= 4:
                        break
            if len(prices) >= 4:
                return prices[:4]
        return None

    # === СТАРЫЕ ФУНКЦИИ ДЛЯ WORD (БОЛЬШЕ НЕ ИСПОЛЬЗУЕМ, НО ОСТАВЛЯЕМ ДЛЯ СОВМЕСТИМОСТИ) ===
    # Они заменены на новые выше, но оставляем чтобы не ломать код

    def extract_price_from_text(text):
        """Извлекает цену из текста"""
        if not text:
            return "❌"

        cleaned = re.sub(r'[^\d\s]', '', str(text))
        cleaned = cleaned.replace(' ', '')

        if cleaned and cleaned.isdigit():
            return int(cleaned)

        return "❌"

    def extract_common_prices_universal(filepath):
        """Универсальное извлечение тарифов 'Общий' и 'Общий плюс' из Word файлов"""
        try:
            # Определяем тип файла и конвертируем при необходимости
            file_ext = os.path.splitext(filepath)[1].lower()

            if file_ext == '.doc':
                # Конвертируем .doc в .docx
                converted_path = convert_doc_to_docx(filepath)
                if not converted_path:
                    return ["❌"] * 14
                filepath = converted_path
                file_ext = '.docx'

            if file_ext != '.docx':
                return ["❌"] * 14

            # Основная логика извлечения
            from docx import Document

            doc = Document(filepath)
            target_keys = ["1+4", "1+9", "1+19", "1+49", "1+99", "1+199", "1+499"]

            common_prices = {key: "❌" for key in target_keys}
            common_plus_prices = {key: "❌" for key in target_keys}

            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]

                    if len(row_text) >= 3:
                        key_cell = row_text[0]
                        common_cell = row_text[1] if len(row_text) > 1 else ""
                        common_plus_cell = row_text[2] if len(row_text) > 2 else ""

                        for key in target_keys:
                            if key in key_cell:
                                if common_prices[key] == "❌":
                                    common_prices[key] = clean_price(common_cell)
                                if common_plus_prices[key] == "❌":
                                    common_plus_prices[key] = clean_price(common_plus_cell)

            common_list = [common_prices[key] for key in target_keys]
            common_plus_list = [common_plus_prices[key] for key in target_keys]

            return common_list + common_plus_list

        except Exception as e:
            return ["❌"] * 14

    def clean_price(price_str):
        """Очищает цену от лишних символов"""
        if not price_str:
            return "❌"
        cleaned = re.sub(r'[^\d\s]', '', price_str)
        cleaned = cleaned.replace(' ', '')
        if cleaned and cleaned.isdigit():
            return int(cleaned)
        return "❌"

    # === ФУНКЦИЯ ДЛЯ СКАЧИВАНИЯ ФАЙЛОВ (ОСТАВЛЯЕМ БЕЗ ИЗМЕНЕНИЙ) ===

    def download_file_by_text(text):
        """Улучшенная функция скачивания файлов по тексту ссылки"""
        try:
            # Ждем полной загрузки страницы
            time.sleep(3)

            # Прокручиваем страницу вниз чтобы увидеть все элементы
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)
            driver.execute_script("window.scrollTo(0, 0);")
            time.sleep(1)

            # Пробуем разные стратегии поиска ссылки
            link = None
            strategies = [
                f"//a[contains(text(), '{text}')]",
                f"//a[contains(., '{text.split('«')[0]}')]",
                "//a[contains(@class, 'link')]",
                f"//*[contains(text(), '{text.split()[0]}')]",
            ]

            for strategy in strategies:
                try:
                    link = wait.until(EC.element_to_be_clickable((By.XPATH, strategy)))
                    if link:
                        break
                except Exception as e:
                    continue

            if not link:
                all_links = driver.find_elements(By.TAG_NAME, "a")
                for l in all_links:
                    try:
                        link_text = l.text
                        if text in link_text or any(word in link_text for word in text.split()[:2]):
                            link = l
                            break
                    except:
                        continue

            if not link:
                return None

            # Получаем URL
            file_url = link.get_attribute('href')

            if not file_url:
                return None

            # Прокручиваем к элементу с отступом
            driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", link)
            time.sleep(2)

            # Выделяем элемент для визуализации
            driver.execute_script("arguments[0].style.border='3px solid red';", link)
            time.sleep(1)

            # Пробуем разные методы клика
            try:
                link.click()
            except:
                try:
                    driver.execute_script("arguments[0].click();", link)
                except:
                    from selenium.webdriver.common.action_chains import ActionChains
                    actions = ActionChains(driver)
                    actions.move_to_element(link).click().perform()

            # Увеличиваем время ожидания скачивания
            time.sleep(15)

            # Ищем скачанный файл
            files = [f for f in os.listdir(DOWNLOAD_DIR)
                    if not f.startswith('.') and not f.startswith('~') and not f.endswith('.crdownload')]
            if files:
                latest_file = max([os.path.join(DOWNLOAD_DIR, f) for f in files], key=os.path.getctime)
                file_size = os.path.getsize(latest_file)

                if file_size > 100:
                    return latest_file
                else:
                    return None
            else:
                return None

        except Exception as e:
            return None

    # === ОСНОВНАЯ ЛОГИКА ПАРСИНГА ===
    try:
        # Создаем Excel файл
        wb = Workbook()
        ws = wb.active
        ws.title = "Тарифы"

        # ОБНОВЛЕННЫЕ ЗАГОЛОВКИ С КОЛОНКАМИ ДЛЯ РЕГРЕССИВНЫХ ШКАЛ
        headers = [
            "Код региона", "Название региона",
            "ИП (УСН)", "ИП (ОСНО)", "ЮЛ (УСН)", "ЮЛ (ОСНО)",
            "Бюджетник плюс", "Бюджетник",
            "1+4", "1+9", "1+19", "1+49", "1+99", "1+199", "1+499",
            "1+4 плюс", "1+9 плюс", "1+19 плюс", "1+49 плюс", "1+99 плюс", "1+199 плюс", "1+499 плюс",
            "Нулевая отчетность",
            "Налоговый представитель Базовый",
            "Зона регрессии",
            "до 199", "200-499", "500-999", "1000-1999", "от 2000",
            "Стартовый онлайн ИП (УСН)", "Стартовый онлайн ИП (ОСНО)",
            "Стартовый онлайн ЮЛ (УСН)", "Стартовый онлайн ЮЛ (ОСНО)"
        ]
        ws.append(headers)

        # Создаем строки для всех регионов
        for region_id, region_name in regions:
            row = [int(region_id), region_name] + ["❌"] * (len(headers) - 2)
            ws.append(row)

        # Переходим на страницу для доступа к ссылкам
        driver.get(BASE_URL.replace("77", "01"))
        wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        time.sleep(5)

        # Скачиваем PDF файлы
        null_pdf = download_file_by_text("Скачать прайс-лист на тарифные планы «Общий Лайт», «Нулевая отчетность», «Кадровые отчеты», «Классический»")
        tax_pdf = download_file_by_text("Скачать прайс-лист для налоговых представителей")
        start_pdf = download_file_by_text("Скачать прайс-лист на тарифный план «Стартовый онлайн»")

        # Извлекаем данные из PDF
        null_prices = extract_all_null_prices(null_pdf) if null_pdf else {}
        tax_rep_prices = extract_all_tax_representative_prices(tax_pdf) if tax_pdf else {}
        start_online_prices = extract_all_start_online_prices(start_pdf) if start_pdf else {}

        # ОПРЕДЕЛЕНИЕ КОЛОНОК ДЛЯ РАЗНЫХ ТИПОВ ДАННЫХ

        NULL_COL = 23
        TAX_BASE_COL = 24
        ZONE_COL = 25

        # МАППИНГ ДЛЯ ПЕРВОЙ ТАБЛИЦЫ (зоны 1,2,3,5,6,7,8,9,11,12)
        REGRESSION_COLS_MAIN = {
            'до_199': 26,
            'от_200_до_499': 27,
            'от_500_до_999': 28,
            'от_1000_до_1999': 29,
            'от_2000': 30
        }

        REGRESSION_COLS_4_10 = {
            'до_349': 26,           # Для зон 4 и 10: до 349
            'от_350_до_599': 27,    # Для зон 4 и 10: от 350 до 599
            'от_600_до_999': 28,    # Для зон 4 и 10: от 600 до 999
            'от_1000': 29           # Для зон 4 и 10: от 1000
        }

        START_COLS = [31, 32, 33, 34]

        for row_idx in range(2, len(regions) + 2):
            region_id_cell = ws.cell(row=row_idx, column=1).value
            if region_id_cell is not None:
                region_id = str(region_id_cell).zfill(2)
                region_name = ws.cell(row=row_idx, column=2).value

                # Нулевая отчетность
                if region_id in null_prices:
                    null_price = null_prices[region_id]
                    ws.cell(row=row_idx, column=NULL_COL).value = null_price

                # Налоговый представитель с регрессивными шкалами
                if region_id in tax_rep_prices:
                    tax_data = tax_rep_prices[region_id]

                    if isinstance(tax_data, dict):
                        if 'base_price' in tax_data and tax_data['base_price'] is not None:
                            ws.cell(row=row_idx, column=TAX_BASE_COL).value = tax_data['base_price']

                        if 'zone' in tax_data and tax_data['zone'] is not None:
                            ws.cell(row=row_idx, column=ZONE_COL).value = tax_data['zone']

                        regression_prices = tax_data.get('regression_prices', {})
                        if regression_prices:
                            # Получаем номер зоны
                            zone_number = tax_data.get('zone')

                            # Выбираем правильный маппинг
                            if zone_number in ['4', '10']:
                                regression_mapping = REGRESSION_COLS_4_10
                            else:
                                regression_mapping = REGRESSION_COLS_MAIN

                            for range_key, col_idx in regression_mapping.items():
                                if range_key in regression_prices and regression_prices[range_key] is not None:
                                    ws.cell(row=row_idx, column=col_idx).value = regression_prices[range_key]
                    else:
                        ws.cell(row=row_idx, column=TAX_BASE_COL).value = tax_data

                # Стартовый онлайн
                if region_id in start_online_prices:
                    prices = start_online_prices[region_id]
                    for i, price in enumerate(prices):
                        if i < len(START_COLS):
                            ws.cell(row=row_idx, column=START_COLS[i]).value = price

        # === ОБРАБОТКА WORD ФАЙЛОВ ДЛЯ РЕГИОНОВ ===
        successful_downloads = 0

        for idx, (region_id, region_name) in enumerate(regions, 1):
            if cancel_flag:
                await message.edit_text("❌ Контур: Парсинг отменен.")
                break

            try:
                region_url = BASE_URL.replace("77", region_id)
                driver.get(region_url)
                wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))
                time.sleep(5)

                # Очищаем папку от старых файлов
                for f in os.listdir(DOWNLOAD_DIR):
                    if f.endswith(('.doc', '.docx')):
                        try:
                            os.remove(os.path.join(DOWNLOAD_DIR, f))
                        except:
                            pass

                # Скачиваем Word файл
                word_file = download_file_by_text("Скачать полный прайс-лист, часть 2")

                if word_file:
                    successful_downloads += 1

                    # === НОВАЯ ЛОГИКА ИЗВЛЕЧЕНИЯ ДАННЫХ ===
                    # Извлекаем все данные одной функцией
                    all_prices = extract_prices_universal(word_file)

                    # Распаковываем результаты (22 значения)
                    # Порядок: [ip_usn, ip_osno, ul_usn, ul_osno, budget_plus, budget,
                    #           1+4, 1+9, 1+19, 1+49, 1+99, 1+199, 1+499,
                    #           1+4_plus, 1+9_plus, 1+19_plus, 1+49_plus, 1+99_plus, 1+199_plus, 1+499_plus]

                    ip_usn = all_prices[0] if len(all_prices) > 0 else "❌"
                    ip_osno = all_prices[1] if len(all_prices) > 1 else "❌"
                    ul_usn = all_prices[2] if len(all_prices) > 2 else "❌"
                    ul_osno = all_prices[3] if len(all_prices) > 3 else "❌"
                    budget_plus = all_prices[4] if len(all_prices) > 4 else "❌"
                    budget = all_prices[5] if len(all_prices) > 5 else "❌"

                    # Общие тарифы (7 значений)
                    common_prices = all_prices[6:13] if len(all_prices) >= 13 else ["❌"] * 7

                    # Общие плюс тарифы (7 значений)
                    common_plus_prices = all_prices[13:20] if len(all_prices) >= 20 else ["❌"] * 7

                    # Обновляем данные в Excel
                    row_idx = idx + 1

                    # Основные тарифы
                    ws.cell(row=row_idx, column=3).value = ip_usn      # ИП (УСН)
                    ws.cell(row=row_idx, column=4).value = ip_osno     # ИП (ОСНО)
                    ws.cell(row=row_idx, column=5).value = ul_usn      # ЮЛ (УСН)
                    ws.cell(row=row_idx, column=6).value = ul_osno     # ЮЛ (ОСНО)
                    ws.cell(row=row_idx, column=7).value = budget_plus # Бюджетник плюс
                    ws.cell(row=row_idx, column=8).value = budget      # Бюджетник

                    # Тарифы Общий (колонки 9-15)
                    for i, price in enumerate(common_prices):
                        if i < 7:
                            ws.cell(row=row_idx, column=9 + i).value = price

                    # Тарифы Общий плюс (колонки 16-22)
                    for i, price in enumerate(common_plus_prices):
                        if i < 7:
                            ws.cell(row=row_idx, column=16 + i).value = price

            except Exception as e:
                pass

            # Обновляем прогресс
            progress = int((idx) / total_regions * 100)
            await message.edit_text(f"🔄 Прогресс: {progress}%")

            # Периодически сохраняем Excel
            if idx % 5 == 0:
                wb.save(FILE_NAME_KONTUR)

        # Финальное сохранение
        wb.save(FILE_NAME_KONTUR)

        # === ОТПРАВКА РЕЗУЛЬТАТА В ЧАТ ===
        if os.path.exists(FILE_NAME_KONTUR):
            logging.info(f"Файл {FILE_NAME_KONTUR} создан, отправляем в чат")
            if cancel_flag:
                comment = f"⚠️ Парсинг Контур был отменен. Файл содержит неполные данные"
                logging.info("Парсинг был отменен, отправляем неполный файл")
            else:
                comment = f"✅ Парсинг Контур завершен успешно"

            await send_file_into_chat(TELEGRAM_CHAT_ID, FILE_NAME_KONTUR, comment)
            logging.info("Файл Контур успешно отправлен в чат")
        else:
            await callback_query.message.answer("❌ Не удалось создать файл с результатами")
            logging.error(f"Файл {FILE_NAME_KONTUR} не найден")

    except Exception as e:
        error_msg = f"❌ Ошибка парсинга: {str(e)}"
        logging.error(f"Ошибка в parse_kontur: {str(e)}", exc_info=True)
        try:
            await callback_query.message.answer(error_msg)
        except Exception as e2:
            logging.error(f"Не удалось отправить сообщение об ошибке: {str(e2)}")

    finally:
        try:
            driver.quit()
        except:
            pass

# Запуск бота
async def main():
    await dp.start_polling(bot)

if __name__ == '__main__':
    asyncio.run(main())
